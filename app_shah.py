import streamlit as st
import hashlib
import base64
from datetime import date
import io
import os
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Page config ---
st.set_page_config(page_title="Letter Crafter", layout="wide")
st.title("📄 Letter Crafter: Recommendation Letter Generator")

# --- Password protection ---
def verify_password(pw: str) -> bool:
    return hashlib.sha256(pw.encode()).hexdigest() == st.secrets.get("password_hash", "")

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    pw = st.text_input("Enter password", type="password")
    if pw and verify_password(pw):
        st.session_state.authenticated = True
        st.rerun()
    elif pw:
        st.error("Incorrect password.")
    st.stop()

# --- OpenAI client ---
client = OpenAI(api_key=st.secrets["openai_api_key"])

# --- Inputs ---
st.subheader("📁 Upload Materials")
uploaded_files = st.file_uploader(
    "Upload CVs, drafts, personal statements, etc.", accept_multiple_files=True
)

st.subheader("👥 Describe Your Relationship")
relationship_text = st.text_area("How do you know the applicant? (1–4 sentences)", height=120)

addressee = st.text_input("Addressee (e.g., Admissions Committee)", "")
salutation = st.text_input("Salutation (e.g., Dear Committee)", "")
if not salutation.strip():
    salutation = "To Whom It May Concern"

letter_date = date.today().strftime("%B %d, %Y")
filename = st.text_input("Output filename (no extension)", value="recommendation_letter")

font_name = st.selectbox("Font", ["Arial", "Times New Roman", "Calibri", "Aptos"], index=0)
font_size = st.selectbox("Font size", [9, 10, 10.5, 11, 11.5, 12], index=3)

# --- File base64 preview ---
def prepare_file_context(files):
    previews = []
    for f in files:
        content = f.read()
        encoded = base64.b64encode(content).decode("utf-8")
        preview = encoded[:500]  # limit to reduce token usage
        previews.append(f"{f.name} (base64 preview):\n{preview}...\n")
    return "\n".join(previews)

# --- Generate letter with GPT-4o ---
def generate_letter(relationship_text, files):
    system_prompt = (
        "You are Letter Crafter, an expert letter writer. You will receive a description of the recommender's "
        "relationship with the applicant and base64 previews of attached files (e.g., CVs, drafts, etc). "
        "Use this information to write the body of a polished recommendation letter. "
        "Do NOT include the date, salutation, or closing. Return only the letter body."
    )

    file_context = prepare_file_context(files)
    user_prompt = (
        f"My relationship to the applicant:\n{relationship_text}\n\n"
        f"Attached files:\n{file_context}\n\n"
        f"Please write a professional recommendation letter body only."
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Error generating letter: {e}")
        return None

# --- Generate button ---
if st.button("✍️ Generate Letter"):
    if not uploaded_files or not relationship_text.strip():
        st.warning("Please upload at least one file and describe your relationship.")
        st.stop()

    letter_body = generate_letter(relationship_text, uploaded_files)
    if letter_body:
        st.session_state.letter_text = letter_body
        st.session_state.addressee = addressee
        st.session_state.salutation = salutation
        st.session_state.date = letter_date
        st.success("✅ Letter body generated.")

# --- Template insertion ---
template_path = os.path.join(os.path.dirname(__file__), "Shah_LOS_template.docx")

def replace_placeholders(doc, replacements):
    date_idx = None

    for idx, p in enumerate(doc.paragraphs):
        for placeholder, replacement in replacements.items():
            if placeholder in p.text:
                # Clear and replace
                p.clear()
                run = p.add_run(replacement)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

                # Track the <<Date>> index only
                if placeholder == "<<Date>>":
                    date_idx = idx

    # Only remove excess empty paragraphs after <<Date>>
    if date_idx is not None:
        i = date_idx + 1
        while i < len(doc.paragraphs) and doc.paragraphs[i].text.strip() == "":
            # Stop if we reach the <<Addressee>> or <<Salutation>> block
            if any(ph in doc.paragraphs[i].text for ph in ["<<Addressee>>", "<<Salutation>>"]):
                break
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)


if "letter_text" in st.session_state and os.path.exists(template_path):
    try:
        doc = Document(template_path)

        replacements = {
            "<<Date>>": st.session_state.date,
            "<<Addressee>>": st.session_state.addressee,
            "<<Salutation>>": st.session_state.salutation,
            "<<Enter text here>>": st.session_state.letter_text
        }

        replace_placeholders(doc, replacements)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Download Letter (DOCX)",
            data=buffer,
            file_name=f"{filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Error formatting letter: {e}")
elif not os.path.exists(template_path):
    st.error("📁 Missing Word template: Shah_LOS_template.docx")
